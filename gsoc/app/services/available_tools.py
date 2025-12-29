import os
import re
import time
import base64
import shutil
import subprocess
import tempfile
import datetime
from zoneinfo import ZoneInfo

import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
import docx

# =========================================================
# CENTRAL PROMPTS (EXACTLY AS PROVIDED)
# =========================================================

SUMMARY_PROMPT = """
You are Resumini, an intelligent AI Resume Summarization Agent that produces professional single-paragraph summaries suitable for recruiters and hiring systems.
Do not use markdown, lists, emojis, or decorative symbols.

Task:
Analyze the resume below and generate a single, coherent, professional paragraph summarizing the candidate’s overall profile.

Guidelines:
- Do not use markdown, emojis, or decorative formatting.
- Detect and include the candidate’s full name if mentioned.
- Mention total experience duration (if inferable), current or most recent role, and area of expertise.
- Highlight 4–6 most relevant skills, tools, or technologies naturally within the sentence.
- Mention educational background or domain focus if applicable.
- Avoid repetitive or filler words like “hardworking” or “motivated.”
- Keep tone objective, factual, and recruiter-friendly.
- Output should be a single paragraph of 100–130 words maximum.
- Do not use bullet points, numbered lists, or formatting symbols.
- The summary must read like a natural executive summary written by a hiring analyst.

Output Format (plain text only):

==========================
CANDIDATE SUMMARY REPORT
==========================
Candidate Name: <Extracted name or "Name not found">
Summary:
<One professional paragraph summarizing the resume content.>
==========================

Resume Text:
\"\"\"{resume_text}\"\"\"
"""


ATS_SCORING_PROMPT = """
You are an advanced ATS (Applicant Tracking System) evaluation agent.
Do not use markdown, emojis, or decorative formatting.

Input: Candidate resume text and a job description.
Task: Evaluate how well the resume matches the job role.

Output must strictly follow this structure:

==========================
ATS SCORING REPORT
==========================
Candidate Name: <Name or Not Found>
Target Role: <Extracted or from job description if possible>
ATS Match Score: <Score out of 100>

Feedback Summary:
Provide a short plain-text explanation in 2–3 sentences describing how the candidate’s resume aligns or misaligns with the target role.
==========================

Resume:
\"\"\"{resume}\"\"\"

Job Description:
\"\"\"{job}\"\"\"
"""


OPTIMIZE_PROMPT = """
You are an expert Resume Optimization Agent.
Do not use markdown, emojis, or decorative symbols in the response.

Input: Candidate resume and target job role.
Task: Rewrite the resume to better align with the target role while preserving authenticity.

Guidelines:
- Keep tone formal, concise, and ATS-friendly.
- Insert relevant keywords naturally without exaggeration.
- Preserve measurable data (years, percentages, metrics).
- Avoid unnecessary adjectives or generic phrasing.
- The final output should be formatted for direct Word document insertion (no headers or separators inside the resume content).

Output must strictly follow this structure:

==========================
OPTIMIZED RESUME REPORT
==========================
Candidate Name: <Name or Not Found>
Target Role: {role}

Optimization Summary:
<2–3 sentence description of improvements made.>

Optimized Resume Content:
<Full optimized resume text here>
==========================

Resume Text:
\"\"\"{resume_text}\"\"\"
"""

# =========================================================
# WEATHER + TIME (SAMPLE TOOLS)
# =========================================================

def get_weather(city: str) -> dict:
    if city.lower() == "new york":
        return {"status": "success", "report": "Sunny, 25°C (77°F)"}
    return {"status": "error", "error_message": f"No weather data for {city}"}


def get_current_time(city: str) -> dict:
    tz_map = {
        "new york": "America/New_York",
        "india": "Asia/Kolkata"
    }
    tz = tz_map.get(city.lower())
    if not tz:
        return {"status": "error", "error_message": "Unknown city"}

    now = datetime.datetime.now(ZoneInfo(tz))
    return {
        "status": "success",
        "report": now.strftime("%Y-%m-%d %H:%M:%S %Z")
    }

# =========================================================
# FILE PARSER
# =========================================================

def extract_text(path: str) -> dict:
    if not os.path.exists(path):
        return {"status": "error", "error_message": "File not found"}

    ext = os.path.splitext(path.lower())[1]

    try:
        if ext == ".pdf":
            import pdfplumber
            pages = []
            with pdfplumber.open(path) as pdf:
                for p in pdf.pages:
                    if p.extract_text():
                        pages.append(p.extract_text())
            text = "\n".join(pages)

        elif ext in [".doc", ".docx"]:
            doc = docx.Document(path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        else:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()

        return {"status": "success", "report": text}

    except Exception as e:
        return {"status": "error", "error_message": str(e)}

# =========================================================
# VECTOR DB (FAISS) – RAG
# =========================================================

_EMBED_MODEL = SentenceTransformer("all-MiniLM-L6-v2")
_VECTOR_DIM = 384
_VECTOR_INDEX = None
_VECTOR_TEXTS = []


def _chunk_text(text: str, size=500, overlap=80):
    words = text.split()
    chunks, i = [], 0
    while i < len(words):
        chunks.append(" ".join(words[i:i+size]))
        i += size - overlap
    return chunks


def rag_store_resume(resume_text: str) -> dict:
    global _VECTOR_INDEX, _VECTOR_TEXTS

    chunks = _chunk_text(resume_text)
    embeddings = _EMBED_MODEL.encode(chunks).astype("float32")

    if _VECTOR_INDEX is None:
        _VECTOR_INDEX = faiss.IndexFlatL2(_VECTOR_DIM)

    _VECTOR_INDEX.add(embeddings)
    _VECTOR_TEXTS.extend(chunks)

    return {
        "status": "success",
        "chunks_added": len(chunks),
        "total_vectors": _VECTOR_INDEX.ntotal
    }


def rag_query(question: str, top_k: int = 4) -> dict:
    if _VECTOR_INDEX is None:
        return {"status": "error", "error_message": "No resume stored"}

    q_emb = _EMBED_MODEL.encode([question]).astype("float32")
    _, idxs = _VECTOR_INDEX.search(q_emb, top_k)

    context = "\n\n".join(_VECTOR_TEXTS[i] for i in idxs[0])

    prompt = f"""
    Answer using ONLY the context below.

    Context:
    {context}

    Question:
    {question}
    """

    answer = rag_query.llm.generate(prompt)

    return {"status": "success", "report": answer.strip()}

# =========================================================
# ATS HEURISTIC REPORT (JSON)
# =========================================================

def ats_report(resume_text: str, role: str) -> dict:
    text = resume_text.lower()

    keywords = ["python", "machine learning", "ai", "sql", "flask", "tensorflow"]
    sections = ["education", "skills", "experience", "projects"]

    keyword_score = len([k for k in keywords if k in text]) / len(keywords) * 100
    structure_score = len([s for s in sections if s in text]) / len(sections) * 100
    length_score = 100 if 400 <= len(text.split()) <= 900 else 70

    overall = round(
        keyword_score * 0.4 +
        structure_score * 0.3 +
        length_score * 0.3, 2
    )

    return {
        "status": "success",
        "role": role,
        "ats_report": {
            "overall_score": overall,
            "keyword_score": round(keyword_score, 2),
            "structure_score": round(structure_score, 2),
            "length_score": length_score
        }
    }

# =========================================================
# ATS AI SCORING (PROMPT-BASED)
# =========================================================

def ats_ai_feedback(resume_text: str, job_description: str) -> dict:
    prompt = ATS_SCORING_PROMPT.format(
        resume=resume_text,
        job=job_description
    )

    response = ats_ai_feedback.llm.generate(prompt)

    score = None
    m = re.search(r"ATS Match Score:\s*(\d+)", response)
    if m:
        score = int(m.group(1))

    return {
        "status": "success",
        "score": score,
        "report": response.strip()
    }

# =========================================================
# SUMMARY TOOL
# =========================================================

def summarize_resume(resume_text: str) -> dict:
    prompt = SUMMARY_PROMPT.format(resume_text=resume_text)
    response = summarize_resume.llm.generate(prompt)
    return {"status": "success", "report": response.strip()}

# =========================================================
# OPTIMIZE RESUME (DOCX)
# =========================================================

def optimize_resume(resume_text: str, role: str) -> dict:
    prompt = OPTIMIZE_PROMPT.format(
        resume_text=resume_text,
        role=role
    )

    response = optimize_resume.llm.generate(prompt)

    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    out_path = os.path.join(desktop, f"optimized_resume_{role}.docx")

    doc = docx.Document()
    for line in response.splitlines():
        doc.add_paragraph(line)
    doc.save(out_path)

    return {
        "status": "success",
        "file_path": out_path,
        "report": response.strip()
    }

# =========================================================
# LaTeX → PDF → HTML PREVIEW
# =========================================================

def latex_resume_preview(resume_text: str, role: str) -> dict:
    prompt = f"""
    Convert the resume into a clean, ATS-safe,
    single-page LaTeX resume for {role}.
    Output ONLY LaTeX.
    """

    latex = latex_resume_preview.llm.generate(prompt).strip()

    if "\\documentclass" not in latex:
        latex = f"""
\\documentclass[11pt]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[T1]{{fontenc}}
\\usepackage{{geometry}}
\\usepackage{{enumitem}}
\\usepackage{{hyperref}}
\\geometry{{margin=0.75in}}
\\pagestyle{{empty}}
\\begin{{document}}
{latex}
\\end{{document}}
"""

    pdf64 = None

    with tempfile.TemporaryDirectory() as tmp:
        tex = os.path.join(tmp, "r.tex")
        pdf = os.path.join(tmp, "r.pdf")
        open(tex, "w", encoding="utf-8").write(latex)

        if shutil.which("pdflatex"):
            subprocess.run(
                ["pdflatex", tex],
                cwd=tmp,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )

        if os.path.exists(pdf):
            pdf64 = base64.b64encode(open(pdf, "rb").read()).decode()

    html = f"""
    <html><body style="background:#000;color:#fff">
    <h2>Optimized Resume ({role})</h2>
    <embed src="data:application/pdf;base64,{pdf64}"
           width="100%" height="600px">
    </body></html>
    """

    return {
        "status": "success",
        "latex_source": latex,
        "pdf_base64": pdf64,
        "html_preview": html
    }

# =========================================================
# LINKEDIN JOB SEARCH (DEMO)
# =========================================================

def search_jobs(role: str, location: str = "India") -> dict:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.chrome.options import Options

    options = Options()
    options.add_argument("--headless=new")
    driver = webdriver.Chrome(options=options)

    jobs = []
    try:
        url = f"https://www.linkedin.com/jobs/search/?keywords={role}&location={location}"
        driver.get(url)
        time.sleep(5)

        cards = driver.find_elements(By.CSS_SELECTOR, "ul.jobs-search__results-list li")
        for c in cards[:10]:
            try:
                jobs.append({
                    "title": c.find_element(By.TAG_NAME, "h3").text,
                    "company": c.find_element(By.TAG_NAME, "h4").text,
                    "link": c.find_element(By.TAG_NAME, "a").get_attribute("href")
                })
            except Exception:
                continue
    finally:
        driver.quit()

    return {"status": "success", "report": jobs}
